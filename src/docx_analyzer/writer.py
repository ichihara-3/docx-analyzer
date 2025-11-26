"""DOCX comment injection module for AI review results."""

from __future__ import annotations

import re
import tempfile
from typing import Dict, List, Optional, Tuple, Union

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .llm_review import CommentLocator


# Type alias for comment target
# str: quoted text
# Tuple[str, str]: (start_text, end_text)
# None: whole paragraph
CommentTarget = Union[str, Tuple[str, str], None]


def parse_review_for_comments(review_text: str) -> Dict[int, List[Tuple[CommentTarget, str]]]:
    """Parse markdown review text to extract paragraph references and comments.
    
    Args:
        review_text: Markdown formatted review text with paragraph references
        
    Returns:
        Dictionary mapping paragraph index to list of (target, comment_text) tuples.
        
    Example:
        Input: "- [段落 5] \"損害賠償\" 上限が不明確"
        Output: {5: [("損害賠償", "上限が不明確")]}
    """
    comments_map: Dict[int, List[Tuple[CommentTarget, str]]] = {}
    
    # Pattern to match: [段落 X] followed by anything
    # We allow for optional markdown bolding/italic markers before the bracket
    pattern = r'^\s*[-*\d.]*\s*(?:[*_]{2})?\s*\[段落\s+(\d+)\]\s*(.+)$'
    
    # Pattern for quoted text: "text"
    quote_pattern = r'^\s*(?:[*_]{2})?\s*"([^"]+)"'
    
    for line in review_text.split('\n'):
        match = re.match(pattern, line.strip())
        if match:
            para_index = int(match.group(1))
            full_text = match.group(2).strip()
            
            comment_text = full_text
            target: CommentTarget = None
            
            # Check for quoted target
            quote_match = re.match(quote_pattern, full_text)
            if quote_match:
                target = quote_match.group(1)
            
            if para_index not in comments_map:
                comments_map[para_index] = []
            comments_map[para_index].append((target, comment_text))
    
    return comments_map


def split_run_at(paragraph: Paragraph, run: Run, split_text: str) -> List[Run]:
    """Split a run into multiple runs based on split_text.
    
    This is a placeholder for future implementation of true run splitting.
    Currently returns the original run.
    """
    return [run]


def find_runs_for_range(paragraph: Paragraph, start_text: str, end_text: str) -> List[Run]:
    """Find runs that cover the text range from start_text to end_text."""
    if not paragraph.runs:
        return []
        
    full_text = paragraph.text
    
    # Find start index
    start_idx = full_text.find(start_text)
    if start_idx == -1:
        return []
        
    # Find end index (searching from start)
    end_idx_rel = full_text[start_idx:].find(end_text)
    if end_idx_rel == -1:
        return []
        
    end_idx = start_idx + end_idx_rel + len(end_text)
    
    # Now map character indices to runs
    target_runs = []
    current_idx = 0
    
    for run in paragraph.runs:
        run_len = len(run.text)
        run_start = current_idx
        run_end = current_idx + run_len
        
        # Check if run overlaps with [start_idx, end_idx)
        if run_end > start_idx and run_start < end_idx:
            target_runs.append(run)
            
        current_idx += run_len
        
    return target_runs


def inject_comments_to_docx(
    source_path: str,
    output_path: str,
    comments_map: Dict[int, List[Tuple[CommentTarget, str]]],
    author: str = "AI Reviewer"
) -> None:
    """Inject comments into a DOCX file at specified paragraphs."""
    doc = Document(source_path)
    
    for para_index, comments in comments_map.items():
        if para_index >= len(doc.paragraphs):
            continue
            
        paragraph = doc.paragraphs[para_index]
        if not paragraph.runs:
            continue
            
        for target, comment_text in comments:
            target_runs = []
            
            if isinstance(target, tuple):
                # Range target: (start, end)
                start_text, end_text = target
                target_runs = find_runs_for_range(paragraph, start_text, end_text)
                
                # Fallback to whole paragraph if range not found
                if not target_runs:
                    target_runs = paragraph.runs
                    
            elif isinstance(target, str):
                # Quoted target
                for run in paragraph.runs:
                    if target in run.text:
                        target_runs = [run]
                        break
                if not target_runs:
                    target_runs = paragraph.runs
            else:
                # Whole paragraph
                target_runs = paragraph.runs
            
            try:
                doc.add_comment(
                    runs=target_runs,
                    text=comment_text,
                    author=author,
                    initials="AI"
                )
            except Exception:
                continue
    
    doc.save(output_path)


def create_commented_docx(
    source_path: str,
    review_text: str,
    author: str = "AI Reviewer"
) -> str:
    """Create a DOCX file with AI review comments injected."""
    # 1. Parse initial review
    comments_map = parse_review_for_comments(review_text)
    
    # 2. Refine locations using CommentLocator
    # We need to read the docx to get paragraph texts
    doc = Document(source_path)
    locator = CommentLocator()
    
    refined_map: Dict[int, List[Tuple[CommentTarget, str]]] = {}
    
    for para_index, comments in comments_map.items():
        refined_map[para_index] = []
        
        if para_index >= len(doc.paragraphs):
            refined_map[para_index] = comments
            continue
            
        paragraph_text = doc.paragraphs[para_index].text
        
        # Identify comments that need location refinement (those with None target)
        comments_to_locate = []
        indices_to_locate = []
        
        for i, (target, text) in enumerate(comments):
            if target is None:
                comments_to_locate.append(text)
                indices_to_locate.append(i)
        
        # Call locator if needed
        locations = []
        if comments_to_locate:
            try:
                locations = locator.locate_comments(paragraph_text, comments_to_locate)
            except Exception:
                # Fallback on error
                locations = [{"start": "", "end": ""} for _ in comments_to_locate]
        
        # Reconstruct the list with refined targets
        loc_idx = 0
        for i, (target, text) in enumerate(comments):
            if i in indices_to_locate:
                loc = locations[loc_idx]
                start = loc.get("start")
                end = loc.get("end")
                if start and end:
                    refined_map[para_index].append(((start, end), text))
                else:
                    refined_map[para_index].append((None, text))
                loc_idx += 1
            else:
                refined_map[para_index].append((target, text))

    # 3. Inject comments
    temp_file = tempfile.NamedTemporaryFile(
        suffix=".docx",
        delete=False,
        prefix="commented_"
    )
    temp_path = temp_file.name
    temp_file.close()
    
    inject_comments_to_docx(source_path, temp_path, refined_map, author)
    
    return temp_path


