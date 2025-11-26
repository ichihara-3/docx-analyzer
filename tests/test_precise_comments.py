import unittest
import tempfile
import os
from docx import Document
from docx_analyzer.writer import parse_review_for_comments, inject_comments_to_docx

class TestPreciseComments(unittest.TestCase):
    def test_parse_review_for_comments(self):
        review_text = """
        - [段落 1] "target" comment 1
        - [段落 2] comment 2
        - [段落 3] "multi word target" comment 3
        """
        expected = {
            1: [("target", '"target" comment 1')],
            2: [(None, "comment 2")],
            3: [("multi word target", '"multi word target" comment 3')]
        }
        result = parse_review_for_comments(review_text)
        self.assertEqual(result, expected)

    def test_inject_comments_to_docx(self):
        # Create a dummy docx
        doc = Document()
        p1 = doc.add_paragraph()
        p1.add_run("This is paragraph 1 with target text.")
        p2 = doc.add_paragraph()
        p2.add_run("This is paragraph 2.")
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            source_path = tmp.name
        doc.save(source_path)
        
        output_path = source_path.replace(".docx", "_out.docx")
        
        comments_map = {
            0: [("target", '"target" comment on target')],
            1: [(None, "comment on paragraph")]
        }
        
        try:
            inject_comments_to_docx(source_path, output_path, comments_map)
            self.assertTrue(os.path.exists(output_path))
            
        finally:
            if os.path.exists(source_path):
                os.remove(source_path)
            if os.path.exists(output_path):
                os.remove(output_path)

if __name__ == '__main__':
    unittest.main()
